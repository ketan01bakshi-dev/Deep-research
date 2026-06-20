"""Generate research reports as Markdown and DOCX."""

from __future__ import annotations

import json
import re
from typing import Any

from cursor_sdk import CustomTool, CustomToolContext
from docx import Document
from docx.shared import Pt

from cursor_agent_core.paths.file_helpers import (
    relative_to_project,
    sanitize_filename,
    slugify,
    unique_path,
    utc_timestamp,
)
from cursor_agent_core.paths.project_context import get_output_directory, get_project_root

_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_BOLD_RE = re.compile(r"\*\*([^*]+)\*\*")


def _parse_sections(raw_sections: Any) -> list[dict[str, str]]:
    if not isinstance(raw_sections, list) or not raw_sections:
        raise ValueError("sections is required (non-empty list of {heading, body})")

    sections: list[dict[str, str]] = []
    for item in raw_sections:
        if not isinstance(item, dict):
            raise ValueError("each section must be an object with heading and body")
        heading = str(item.get("heading") or "").strip()
        body = str(item.get("body") or "").strip()
        if not heading:
            raise ValueError("each section requires a heading")
        sections.append({"heading": heading, "body": body})
    return sections


def _parse_sources(raw_sources: Any) -> list[dict[str, str]]:
    if raw_sources is None:
        return []
    if not isinstance(raw_sources, list):
        raise ValueError("sources must be a list of {title, url} objects")

    sources: list[dict[str, str]] = []
    for item in raw_sources:
        if not isinstance(item, dict):
            continue
        title = str(item.get("title") or "").strip()
        url = str(item.get("url") or "").strip()
        if title or url:
            sources.append({"title": title or url, "url": url})
    return sources


def _parse_cited_papers(raw_papers: Any) -> list[str]:
    if raw_papers is None:
        return []
    if not isinstance(raw_papers, list):
        raise ValueError("cited_papers must be a list of project-relative file paths")
    return [str(p).strip() for p in raw_papers if str(p).strip()]


def _markdown_report(
    *,
    title: str,
    sections: list[dict[str, str]],
    sources: list[dict[str, str]],
    key_takeaways: list[str],
    disclaimer: str,
    cited_papers: list[str],
) -> str:
    lines = [
        "---",
        f'title: "{title}"',
        f"date: {utc_timestamp()}",
        "---",
        "",
        f"# {title}",
        "",
    ]

    for section in sections:
        lines.append(f"## {section['heading']}")
        lines.append("")
        lines.append(section["body"])
        lines.append("")

    if key_takeaways:
        lines.append("## Key Takeaways")
        lines.append("")
        for item in key_takeaways:
            lines.append(f"- {item}")
        lines.append("")

    if sources:
        lines.append("## Sources")
        lines.append("")
        for index, source in enumerate(sources, start=1):
            if source["url"]:
                lines.append(f"{index}. [{source['title']}]({source['url']})")
            else:
                lines.append(f"{index}. {source['title']}")
        lines.append("")

    if cited_papers:
        lines.append("## Downloaded Papers")
        lines.append("")
        for paper in cited_papers:
            lines.append(f"- {paper}")
        lines.append("")

    if disclaimer:
        lines.append("## Disclaimer")
        lines.append("")
        lines.append(disclaimer)
        lines.append("")

    return "\n".join(lines)


def _add_markdown_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    cursor = 0
    for match in _LINK_RE.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        label = match.group(1)
        url = match.group(2)
        link_run = paragraph.add_run(label)
        link_run.underline = True
        paragraph.add_run(f" ({url})")
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _add_body(document: Document, body: str) -> None:
    for line in body.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("- "):
            document.add_paragraph(stripped[2:], style="List Bullet")
            continue
        if stripped.startswith("## "):
            document.add_heading(stripped[3:], level=2)
            continue
        if stripped.startswith("### "):
            document.add_heading(stripped[4:], level=3)
            continue
        bold_match = _BOLD_RE.fullmatch(stripped)
        if bold_match:
            run = document.add_paragraph().add_run(bold_match.group(1))
            run.bold = True
            continue
        _add_markdown_paragraph(document, stripped)


def _build_docx(
    *,
    title: str,
    sections: list[dict[str, str]],
    sources: list[dict[str, str]],
    key_takeaways: list[str],
    disclaimer: str,
    cited_papers: list[str],
    dest: Any,
) -> None:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    document.add_heading(title, level=0)
    meta = document.add_paragraph(f"Generated: {utc_timestamp()}")
    meta.style = "Intense Quote"

    for section in sections:
        document.add_heading(section["heading"], level=1)
        _add_body(document, section["body"])

    if key_takeaways:
        document.add_heading("Key Takeaways", level=1)
        for item in key_takeaways:
            document.add_paragraph(item, style="List Bullet")

    if sources:
        document.add_heading("Sources", level=1)
        for index, source in enumerate(sources, start=1):
            line = f"{index}. {source['title']}"
            if source["url"]:
                line += f" — {source['url']}"
            document.add_paragraph(line)

    if cited_papers:
        document.add_heading("Downloaded Papers", level=1)
        for paper in cited_papers:
            document.add_paragraph(paper, style="List Bullet")

    if disclaimer:
        document.add_heading("Disclaimer", level=1)
        document.add_paragraph(disclaimer)

    document.save(dest)


def create_research_report(
    args: dict[str, Any],
    _ctx: CustomToolContext,
) -> str:
    """Create a Markdown and DOCX research report in the session Reports/ folder."""
    title = str(args.get("title") or "").strip()
    if not title:
        raise ValueError("title is required")

    sections = _parse_sections(args.get("sections"))
    sources = _parse_sources(args.get("sources"))
    cited_papers = _parse_cited_papers(args.get("cited_papers"))

    raw_takeaways = args.get("key_takeaways") or []
    key_takeaways = (
        [str(item).strip() for item in raw_takeaways if str(item).strip()]
        if isinstance(raw_takeaways, list)
        else []
    )

    disclaimer = str(args.get("disclaimer") or "").strip()
    if not disclaimer:
        disclaimer = (
            "This report provides educational information only and is not "
            "personal medical advice. Consult a qualified healthcare professional "
            "for individual guidance."
        )

    directory = get_output_directory('reports')
    base_name = slugify(title)
    md_path = unique_path(directory, f"{base_name}.md")
    docx_path = md_path.with_suffix(".docx")

    markdown = _markdown_report(
        title=title,
        sections=sections,
        sources=sources,
        key_takeaways=key_takeaways,
        disclaimer=disclaimer,
        cited_papers=cited_papers,
    )
    md_path.write_text(markdown, encoding="utf-8")
    _build_docx(
        title=title,
        sections=sections,
        sources=sources,
        key_takeaways=key_takeaways,
        disclaimer=disclaimer,
        cited_papers=cited_papers,
        dest=docx_path,
    )

    payload = {
        "title": title,
        "reports_directory": relative_to_project(directory, get_project_root()),
        "markdown_path": relative_to_project(md_path, get_project_root()),
        "docx_path": relative_to_project(docx_path, get_project_root()),
        "section_count": len(sections),
        "source_count": len(sources),
    }
    return json.dumps(payload, indent=2)


CREATE_REPORT_TOOL = CustomTool(
    execute=create_research_report,
    description=(
        "Create a formal research report as Markdown and DOCX in the session Reports/ folder. "
        "Pass structured sections, sources, key takeaways, and optional cited paper paths."
    ),
    input_schema={
        "type": "object",
        "properties": {
            "title": {"type": "string", "description": "Report title."},
            "sections": {
                "type": "array",
                "description": "Report sections with heading and markdown body.",
                "items": {
                    "type": "object",
                    "properties": {
                        "heading": {"type": "string"},
                        "body": {"type": "string"},
                    },
                    "required": ["heading", "body"],
                },
            },
            "sources": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "title": {"type": "string"},
                        "url": {"type": "string"},
                    },
                },
            },
            "key_takeaways": {
                "type": "array",
                "items": {"type": "string"},
            },
            "cited_papers": {
                "type": "array",
                "items": {"type": "string"},
                "description": "Project-relative paths to downloaded papers.",
            },
            "disclaimer": {"type": "string"},
        },
        "required": ["title", "sections"],
    },
)
