"""Shared helpers for generating Deep Research Word documents."""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_issue(
    doc: Document,
    title: str,
    symptom: str,
    root_cause: str,
    fix: str,
    verification: str,
    notes: str = "",
) -> None:
    add_heading(doc, title, level=2)
    add_para(doc, "Symptom", bold=True)
    add_para(doc, symptom)
    add_para(doc, "Root cause", bold=True)
    add_para(doc, root_cause)
    add_para(doc, "Fix applied", bold=True)
    add_para(doc, fix)
    add_para(doc, "Verification", bold=True)
    add_para(doc, verification)
    if notes:
        add_para(doc, "Additional notes", bold=True)
        add_para(doc, notes)
    doc.add_paragraph()


def add_title_page(doc: Document, title: str, subtitle: str) -> None:
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(subtitle)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row_data in rows:
        row = table.add_row().cells
        for idx, cell in enumerate(row_data):
            row[idx].text = cell
    doc.add_paragraph()
