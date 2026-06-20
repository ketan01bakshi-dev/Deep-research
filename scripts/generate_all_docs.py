"""Generate all Word documentation for the Deep Research agent."""

from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

from docx import Document

from doc_common import (
    add_bullets,
    add_heading,
    add_issue,
    add_numbered,
    add_para,
    add_table,
    add_title_page,
)

ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"
AGENTS_DOCS = ROOT.parent / "docs"
TODAY = date.today().isoformat()


def build_master_guide() -> Path:
    out = DOCS / "Deep_Research_Agent_Master_Guide.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    add_title_page(
        doc,
        "Deep Research Agent — Master Guide",
        f"Complete operator and developer reference — {TODAY}",
    )

    add_heading(doc, "1. What Is Deep Research?", 1)
    add_para(
        doc,
        "Deep Research is an autonomous Cursor SDK agent that runs structured research missions "
        "across any field. It searches the web (Tavily), synthesizes answers, and optionally "
        "produces artifacts: formal reports (Markdown + DOCX), mind-map diagrams, PowerPoint "
        "slide decks, and downloaded PDF papers. A Streamlit web UI provides session history, "
        "PDF viewing, session linking, and multi-turn follow-up.",
    )

    add_heading(doc, "2. Agent Phases", 1)
    add_table(
        doc,
        ["Phase", "Entry point", "Module", "Behavior"],
        [
            ("1 — Stateless", "run_research.py", "agent.py", "One-shot query, no memory"),
            ("2 — Conversational", "run_conversation.py", "conversation.py", "Multi-turn with agent resume"),
            ("3 — Autonomous", "run_autonomous.py", "autonomous.py", "Mission-driven, no mid-run questions"),
            ("Web UI", "run_app.cmd / app.py", "deep_research/web/", "Streamlit session browser + form"),
        ],
    )

    add_heading(doc, "3. Project Layout", 1)
    add_bullets(
        doc,
        [
            "deep_research/ — Python package (domain logic, web UI, tools)",
            "packages/cursor_agent_core (parent) — shared SDK bridge, outputs, sessions",
            "packages/agent_core (parent) — shared .env loading",
            "Docs/Research/{session_id}/ — per-query session folders (gitignored)",
            "Docs/Papers/, Reports/, Diagrams/, Slides/ — legacy flat dirs (CLI fallback)",
            "missions/ — example mission JSON files",
            "app.py — Streamlit entry point",
            "scripts/ — doc generation, setup helpers",
            ".cursor/rules/deep-research.mdc — Cursor IDE rules",
        ],
    )

    add_heading(doc, "4. Quick Start", 1)
    add_numbered(
        doc,
        [
            "Run D:\\Agents\\install_windows_dev.cmd (Windows Node/npm fix).",
            "cd D:\\Agents\\Deep research",
            "python -m venv .venv",
            ".venv\\Scripts\\pip install -e ..\\packages\\cursor_agent_core -e ..\\packages\\agent_core",
            ".venv\\Scripts\\pip install -r requirements.txt",
            "copy .env.example to .env — set CURSOR_API_KEY and TAVILY_API_KEY",
            "Optional: setup_diagrams.cmd for Mermaid mind maps",
            "CLI: .venv\\Scripts\\python run_autonomous.py --example protein_digestion_full",
            "Web: .\\run_app.cmd",
        ],
    )

    add_heading(doc, "5. Deliverables Model", 1)
    add_para(
        doc,
        "Users must specify deliverables before a run. If none are selected, the agent does not "
        "start and returns a clarification message. Artifacts are created only for checked items.",
    )
    add_table(
        doc,
        ["Deliverable", "Output", "Tool"],
        [
            ("answer", "Text synthesis in response", "(agent response)"),
            ("report", "Markdown + DOCX", "create_research_report"),
            ("mindmap", "PNG/SVG mind map", "create_diagram (diagram_type=mindmap)"),
            ("slides", "PowerPoint .pptx", "create_slide_deck"),
            ("papers", "PDF files", "download_research_pdfs (+ Tavily search)"),
        ],
    )

    add_heading(doc, "6. Session Folders", 1)
    add_para(doc, "Each research query creates:", bold=True)
    add_bullets(
        doc,
        [
            "Docs/Research/{YYYYMMDD-HHMMSS}_{slug}_{uuid}/",
            "meta.json — topic, status, deliverables, links, agent_id",
            "turns.json — multi-turn thread history",
            "activity.json — tool/status log for UI",
            "completion.json — agent mission manifest",
            "Papers/, Reports/, Diagrams/, Slides/ — session artifacts",
        ],
    )

    add_heading(doc, "7. Documentation Set", 1)
    add_bullets(
        doc,
        [
            "Deep_Research_Agent_Master_Guide.docx (this document)",
            "Deep_Research_Agent_Master_Troubleshooting_Guide.docx",
            "Deep_Research_Agent_Architecture.docx",
            "Deep_Research_Agent_CLI_Reference.docx",
            "Deep_Research_Streamlit_UI_Reference.docx",
            "Deep_Research_Custom_Tools_Reference.docx",
            "Deep_Research_Session_Management_Guide.docx",
        ],
    )
    add_para(doc, "Regenerate: .venv\\Scripts\\python scripts\\generate_all_docs.py", bold=True)

    add_heading(doc, "8. Relationship to D:\\Agents Workspace", 1)
    add_para(
        doc,
        "Deep Research is a Cursor SDK agent family member. It uses cursor_agent_core (not agent_core "
        "for pipeline logic). Job Search and Grocery List use agent_core for CLI batch pipelines. "
        "See D:\\Agents\\docs\\ARCHITECTURE.md and Agents_Folder_Master_Guide.docx.",
    )

    add_heading(doc, "9. Development History (Phases 0–3 + UI)", 1)
    add_table(
        doc,
        ["Milestone", "What changed"],
        [
            ("Phase 0 — Claude SDK", "Initial stateless agent with claude-agent-sdk query() and WebSearch"),
            ("Phase 1 — Stateless research", "run_research.py, health examples (protein digestion)"),
            ("Phase 2 — Conversational memory", "ClaudeSDKClient, run_conversation.py, resume by session ID"),
            ("Cursor migration", "Replaced Claude SDK with cursor-sdk; CURSOR_API_KEY; Tavily for web search"),
            ("Phase 3 — Autonomous missions", "ResearchMission schema, deliverables gate, complete_mission tool"),
            ("Custom tools", "Reports, diagrams, slides, PDF download; Mermaid mind maps via mmdc"),
            ("Streamlit UI", "app.py, session browser, PDF viewer, per-query Docs/Research/ folders"),
            ("Session management", "Linking, follow-up turns, pin/delete/export/retry/branch, activity log"),
            ("Monorepo refactor", "Extracted cursor_agent_core + agent_core.load_agent_env; workspace docs"),
        ],
    )

    add_heading(doc, "10. Authentication", 1)
    add_para(
        doc,
        "Set CURSOR_API_KEY in D:\\Agents\\Deep research\\.env (copy from .env.example). "
        "The key is loaded at startup by agent_core.load_agent_env(). For Tavily web search and "
        "PDF tools, also set TAVILY_API_KEY. The project originally used ANTHROPIC_API_KEY under "
        "Claude SDK; that was replaced during the Cursor migration.",
    )

    add_heading(doc, "11. Cursor IDE Rules", 1)
    add_para(
        doc,
        "Agent behavior for development in Cursor IDE is guided by "
        ".cursor/rules/deep-research.mdc — deliverables model, artifact paths, and mission conventions.",
    )

    doc.save(out)
    return out


def build_troubleshooting_guide() -> Path:
    out = DOCS / "Deep_Research_Agent_Master_Troubleshooting_Guide.docx"
    doc = Document()
    add_title_page(
        doc,
        "Deep Research Agent — Master Troubleshooting Guide",
        f"Consolidated issues and fixes — {TODAY}",
    )

    add_heading(doc, "1. Executive Summary", 1)
    add_para(
        doc,
        "This guide consolidates issues encountered building and operating Deep Research: "
        "Cursor SDK on Windows, Tavily search, Mermaid diagrams, Streamlit UI, session folders, "
        "and the D:\\Agents monorepo refactor (cursor_agent_core).",
    )
    add_bullets(
        doc,
        [
            "Always use .venv\\Scripts\\python.exe — not system Python.",
            "CURSOR_API_KEY must be in .env; Streamlit must load it via agent_core.load_agent_env.",
            "TAVILY_API_KEY required for live web search and PDF download tools.",
            "After code changes to cursor_agent_core: pip install -e ..\\packages\\cursor_agent_core",
            "Restart Streamlit after installing packages (stale process caches old modules).",
        ],
    )

    add_heading(doc, "2. Environment Checklist", 1)
    add_table(
        doc,
        ["Check", "Command"],
        [
            ("Python venv", ".venv\\Scripts\\python.exe --version"),
            ("cursor_agent_core", ".venv\\Scripts\\python -c \"import cursor_agent_core\""),
            ("deep_research", ".venv\\Scripts\\python -c \"import deep_research\""),
            ("CURSOR_API_KEY", ".venv\\Scripts\\python -c \"import os; print(bool(os.getenv('CURSOR_API_KEY')))\""),
            ("TAVILY_API_KEY", ".venv\\Scripts\\python -c \"import os; print(bool(os.getenv('TAVILY_API_KEY')))\""),
            ("Node.js", "node --version"),
            ("npm/npx", "npx.cmd --version (Windows)"),
        ],
    )

    add_heading(doc, "3. Issue Log", 1)

    add_issue(
        doc,
        "Issue 3.1 — CURSOR_API_KEY not set (Streamlit sessions show error)",
        symptom="All research sessions fail immediately with status 'error'. meta.json error: "
        "'Agent.create requires api_key'. UI shows red error badge, no answer stored.",
        root_cause="app.py did not call load_dotenv() before running the agent. Streamlit "
        "process does not inherit .env automatically.",
        fix="app.py now calls agent_core.load_agent_env(PROJECT_ROOT) at startup. Copy "
        ".env.example to .env with valid CURSOR_API_KEY. Restart run_app.cmd.",
        verification="New session completes with status 'completed' and answer in meta.json.",
        notes="Existing error sessions remain in sidebar; use Delete or Clear failed.",
    )
    add_issue(
        doc,
        "Issue 3.2 — ImportError: run_follow_up_mission",
        symptom="Streamlit crashes on load: cannot import name 'run_follow_up_mission' "
        "from deep_research.autonomous.",
        root_cause="Streamlit process started before autonomous.py was saved, or stale "
        "__pycache__ from partial update.",
        fix="Restart Streamlit. Clear __pycache__ if needed. Verify: "
        ".venv\\Scripts\\python -c \"from deep_research.autonomous import run_follow_up_mission\"",
        verification="app.py imports without error; follow-up form works on session detail page.",
    )
    add_issue(
        doc,
        "Issue 3.3 — PowerShell blocks npx / npm (execution policy)",
        symptom="Diagram rendering fails. npx.ps1 cannot be loaded because running scripts "
        "is disabled. Mermaid CLI never runs.",
        root_cause="Windows PowerShell execution policy blocks .ps1 shims.",
        fix="Run D:\\Agents\\install_windows_dev.cmd. cursor_agent_core.node.cli uses "
        "npx.cmd directly. For local mmdc: setup_diagrams.cmd.",
        verification="npx.cmd --version works. create_diagram produces PNG in session Diagrams/.",
    )
    add_issue(
        doc,
        "Issue 3.4 — Puppeteer / Chrome cache corrupt",
        symptom="mmdc fails with Puppeteer browser launch errors.",
        root_cause="Corrupt .puppeteer-cache in project root.",
        fix="Delete Deep research\\.puppeteer-cache. Re-run setup_diagrams.cmd "
        "(puppeteer install chrome-headless-shell).",
        verification="Mind map PNG renders in Docs/Research/.../Diagrams/.",
    )
    add_issue(
        doc,
        "Issue 3.5 — No artifacts despite deliverables selected",
        symptom="Mission completes but report/mindmap/slides/papers tabs empty.",
        root_cause="Agent did not call output tools, or deliverables not in mission prompt. "
        "Autonomous agent may retry once for missing artifact deliverables.",
        fix="Check activity.json for tool calls. Increase --max-retries on CLI. "
        "Verify TAVILY_API_KEY for papers. Re-run with Retry on session detail page.",
        verification="verify_deliverables() returns empty; artifacts appear in session subfolders.",
    )
    add_issue(
        doc,
        "Issue 3.6 — Bridge launch failed (OSError)",
        symptom="error_during_execution with 'Bridge launch failed'.",
        root_cause="Cursor SDK bridge cannot start — wrong cwd, Cursor not installed, "
        "or Windows pipe discovery failure.",
        fix="Ensure Cursor IDE/SDK bridge available. windows_compat patch applied via "
        "cursor_agent_core. Run from Deep research project root (stable cwd for resume).",
        verification="run_autonomous.py completes at least one turn without bridge error.",
    )
    add_issue(
        doc,
        "Issue 3.7 — Invalid API key (Claude era / placeholder .env)",
        symptom="Terminal shows 'Invalid API key' or 'Claude Code returned an error result: success'. "
        "Research fails before any answer is produced.",
        root_cause="ANTHROPIC_API_KEY missing, placeholder value, or .env reset by copying .env.example.",
        fix="After Cursor migration: use CURSOR_API_KEY instead. Edit .env directly — do not overwrite "
        "with copy .env.example .env if a real key is already set.",
        verification="Agent turn completes; cost/turn count printed in CLI or Streamlit status.",
        notes="Misleading 'success' subtype on auth failure was a Claude SDK quirk.",
    )
    add_issue(
        doc,
        "Issue 3.8 — update_session(pinned=False) ignored",
        symptom="Unpinning a session in Streamlit sidebar does not remove pin from meta.json.",
        root_cause="update_session() only wrote fields when truthy; False was skipped.",
        fix="update_session() now sets all explicitly passed fields including False values.",
        verification="Unpin persists after page refresh; session sorts below pinned items.",
    )
    add_issue(
        doc,
        "Issue 3.9 — pytest not found in venv",
        symptom="python -m pytest fails with 'No module named pytest'.",
        root_cause="pytest not installed in Deep research .venv (dev dependency).",
        fix="pip install pytest in .venv. Run: .venv\\Scripts\\python -m pytest ..\\packages\\cursor_agent_core\\tests",
        verification="2 tests in cursor_agent_core/tests pass.",
    )
    add_issue(
        doc,
        "Issue 3.10 — Linked session context too large",
        symptom="Agent slow or fails when many sessions linked.",
        root_cause="Prior context bundle exceeds token limits (capped at 3 linked sessions).",
        fix="Link at most 3 completed sessions. Truncation applied in get_session_context_bundle().",
        verification="Branch from one session works; prompt includes prior answer summary.",
    )

    add_heading(doc, "4. Quick Command Reference", 1)
    add_table(
        doc,
        ["Task", "Command"],
        [
            ("Setup venv + packages", "pip install -e ..\\packages\\cursor_agent_core -e ..\\packages\\agent_core -r requirements.txt"),
            ("Windows dev setup", "D:\\Agents\\install_windows_dev.cmd"),
            ("Diagram tooling", "setup_diagrams.cmd"),
            ("Stateless research", ".venv\\Scripts\\python run_research.py --example protein_digestion"),
            ("Conversational", ".venv\\Scripts\\python run_conversation.py"),
            ("Autonomous mission", ".venv\\Scripts\\python run_autonomous.py --example protein_digestion_full"),
            ("List mission examples", ".venv\\Scripts\\python run_autonomous.py --list-examples"),
            ("Streamlit UI", "run_app.cmd"),
            ("Regenerate Word docs", ".venv\\Scripts\\python scripts\\generate_all_docs.py"),
            ("cursor_agent_core tests", ".venv\\Scripts\\python -m pytest ..\\packages\\cursor_agent_core\\tests"),
        ],
    )

    add_heading(doc, "5. Known Limitations", 1)
    add_bullets(
        doc,
        [
            "Cancel during run is cooperative — may not stop mid-tool-call instantly.",
            "Agent resume (cursor_agent_id) may expire; follow-up falls back to context reconstruction.",
            "Tavily search quality depends on API key tier and query formulation.",
            "PDF download only works for direct .pdf URLs returned by search.",
            "Health topics include educational disclaimer, not personal medical advice.",
            "Each agent keeps its own .venv — reinstall cursor_agent_core after package changes.",
        ],
    )

    doc.save(out)
    return out


def build_architecture_guide() -> Path:
    out = DOCS / "Deep_Research_Agent_Architecture.docx"
    doc = Document()
    add_title_page(
        doc,
        "Deep Research Agent — Architecture & cursor_agent_core",
        f"System design — {TODAY}",
    )

    add_heading(doc, "1. Design Principle", 1)
    add_para(
        doc,
        "cursor_agent_core answers HOW SDK agents run (bridge, adapter, paths, outputs). "
        "deep_research answers WHAT this agent does (missions, prompts, Tavily, Streamlit UI). "
        "Domain code never lives in cursor_agent_core; shared infra never imports deep_research.",
    )

    add_heading(doc, "2. Dependency Graph", 1)
    add_bullets(
        doc,
        [
            "app.py / run_*.py → deep_research → cursor_agent_core → cursor-sdk",
            "deep_research.tools → cursor_agent_core.outputs + domain tools (Tavily, complete_mission)",
            "session_context.py → cursor_agent_core.sessions.SessionStore (delete, export)",
            "Job Search / Grocery List → agent_core (separate family, no cursor-sdk)",
        ],
    )

    add_heading(doc, "3. cursor_agent_core Modules", 1)
    add_table(
        doc,
        ["Module", "Purpose"],
        [
            ("bridge/windows_compat", "Fix Cursor SDK pipe discovery on Windows"),
            ("bridge/client", "CursorClient.launch_bridge context manager"),
            ("runtime/adapter", "run_agent_turn, run_resume_turn, streaming handlers"),
            ("runtime/options", "build_agent_options with injectable custom_tools"),
            ("runtime/types", "AgentResult dataclass"),
            ("paths/file_helpers", "slugify, unique_path, resolve_under_root"),
            ("paths/project_context", "set_project_root, get_output_directory resolver"),
            ("node/cli", "npm/npx/mmdc resolution on Windows"),
            ("outputs/", "reports, diagrams, slides, pdf_download CustomTools"),
            ("sessions/store", "SessionStore CRUD, turns, activity, zip export"),
        ],
    )

    add_heading(doc, "4. deep_research Modules", 1)
    add_table(
        doc,
        ["Module", "Purpose"],
        [
            ("autonomous.py", "Mission execution, follow-up, session lifecycle"),
            ("autonomous_mission.py", "ResearchMission schema, validation, prompt builder"),
            ("session_context.py", "ResearchSession, linking, turns, activity"),
            ("agent.py", "Phase 1 stateless wrapper"),
            ("conversation.py", "Phase 2 multi-turn agent"),
            ("tools/tavily_search.py", "Literature search custom tool"),
            ("tools/complete_mission.py", "Mission completion manifest"),
            ("tools/paths.py", "Research dirs + session-aware active_* paths"),
            ("web/", "Streamlit UI (browser, form, PDF viewer, follow-up)"),
        ],
    )

    add_heading(doc, "5. Session-Aware Artifact Paths", 1)
    add_para(
        doc,
        "When a research session is active (contextvars), all output tools write to "
        "Docs/Research/{session_id}/Papers|Reports|Diagrams|Slides/. When no active "
        "session (legacy CLI), tools fall back to flat Docs/Papers/, etc.",
    )

    add_heading(doc, "6. Autonomous Mission Flow", 1)
    add_numbered(
        doc,
        [
            "validate_mission() — clarification gate if deliverables empty",
            "create_session() — folder tree + meta.json under Docs/Research/",
            "set_active_session() — contextvar for tool path resolution",
            "build_mission_prompt() — objectives, deliverables, prior linked context",
            "Agent.create() + run_agent_turn() — Cursor SDK execution",
            "verify_deliverables() — optional retry for missing artifacts",
            "update_session() + append_turn() — persist answer, agent_id, status",
            "clear_active_session() in finally block",
        ],
    )

    add_heading(doc, "7. Shim Layer (Backward Compatibility)", 1)
    add_para(
        doc,
        "After extracting cursor_agent_core, deep_research keeps thin shim modules that re-export "
        "shared functionality so existing imports continue to work:",
    )
    add_bullets(
        doc,
        [
            "windows_compat.py → cursor_agent_core.bridge.windows_compat",
            "cursor_client.py → cursor_agent_core.bridge.client",
            "cursor_adapter.py → cursor_agent_core.runtime.adapter",
            "types.py → cursor_agent_core.runtime.types (+ AutonomousResearchResult)",
            "console.py → cursor_agent_core.runtime.console",
            "tools/reports.py, diagrams.py, slides.py, pdf_download.py, node_cli.py → cursor_agent_core",
        ],
    )

    add_heading(doc, "8. Monorepo Integration", 1)
    add_para(
        doc,
        "Deep Research was refactored into D:\\Agents monorepo (2026): shared packages at "
        "D:\\Agents\\packages\\, documentation at D:\\Agents\\docs\\ and Deep research\\docs\\.",
    )
    add_bullets(
        doc,
        [
            "P0: README, ARCHITECTURE.md, AGENT_LAYOUT.md, MODULE_INVENTORY.md",
            "P1: cursor_agent_core bridge + runtime extracted from deep_research",
            "P2: paths + node_cli shared; agent_core.load_agent_env()",
            "P3: output tools (reports, diagrams, slides, pdf) in cursor_agent_core",
            "P4: SessionStore generic CRUD in cursor_agent_core",
            "P5: pyproject.toml workspace, CURSOR_SDK_AGENTS.md scaffold guide",
        ],
    )

    doc.save(out)
    return out


def build_cli_reference() -> Path:
    out = DOCS / "Deep_Research_Agent_CLI_Reference.docx"
    doc = Document()
    add_title_page(
        doc,
        "Deep Research Agent — CLI Reference",
        f"Command-line operator guide — {TODAY}",
    )

    add_heading(doc, "1. Prerequisites", 1)
    add_numbered(
        doc,
        [
            "Python 3.10+ with .venv in project root.",
            "CURSOR_API_KEY in .env (required for all phases).",
            "TAVILY_API_KEY in .env (required for search + PDF tools).",
            "Editable installs: cursor_agent_core, agent_core.",
            "Optional: Node.js + setup_diagrams.cmd for mind maps.",
        ],
    )

    add_heading(doc, "2. SDK Evolution", 1)
    add_para(
        doc,
        "The project began with Claude Agent SDK (query() + ClaudeSDKClient). It was migrated to "
        "Cursor SDK (Agent.create, Agent.prompt, resume). CLI entry points and module names were "
        "preserved where possible. Web search moved from built-in WebSearch to Tavily custom tools.",
    )

    add_heading(doc, "3. Phase 1 — Stateless (run_research.py)", 1)
    add_bullets(
        doc,
        [
            "One-shot research via Agent.prompt(); no session memory.",
            "Uses RESEARCH_SYSTEM_PROMPT embedded in prompt prefix.",
            "Flags: --example, --prompt, --list-examples, --model, --quiet.",
            "Demo examples in deep_research/examples.py (health topics).",
        ],
    )

    add_heading(doc, "4. Phase 2 — Conversational (run_conversation.py)", 1)
    add_bullets(
        doc,
        [
            "Multi-turn research in one Cursor agent session.",
            "Agent.create() persists context; ask() sends follow-up prompts.",
            "resume_research(agent_id, prompt) for cross-process resume.",
            "Examples in conversation_examples.py.",
        ],
    )

    add_heading(doc, "5. Phase 3 — Autonomous (run_autonomous.py)", 1)
    add_bullets(
        doc,
        [
            "Structured ResearchMission: topic, field, objectives, deliverables.",
            "No mid-run questions — agent makes assumptions and lists them.",
            "Creates per-query session folder automatically.",
            "Flags: --example, --mission-file, --list-examples, --max-retries, --quiet.",
            "Exit code 2 on error; 0 on success or clarification.",
        ],
    )

    add_heading(doc, "6. Mission JSON Format", 1)
    add_para(doc, "Example missions/protein_digestion.json:", bold=True)
    add_bullets(
        doc,
        [
            "topic (required), field, objectives (non-empty list)",
            "deliverables: answer | report | mindmap | slides | papers",
            "audience, depth (overview|deep), scope, constraints, success_criteria",
        ],
    )

    add_heading(doc, "7. Mission Examples", 1)
    add_table(
        doc,
        ["Example ID", "Purpose"],
        [
            ("protein_digestion_full", "Full deliverables demo"),
            ("protein_digestion_answer_only", "Text answer only"),
            ("topic_only", "Triggers clarification gate (no deliverables)"),
            ("solid_state_batteries", "Materials science topic"),
        ],
    )

    add_heading(doc, "8. Environment Variables", 1)
    add_table(
        doc,
        ["Variable", "Required", "Purpose"],
        [
            ("CURSOR_API_KEY", "Yes", "Cursor SDK agent authentication"),
            ("TAVILY_API_KEY", "For search", "Tavily literature search + PDF URLs"),
            ("RESEARCH_MODEL", "No", "Override default model (composer-2.5)"),
        ],
    )

    doc.save(out)
    return out


def build_streamlit_reference() -> Path:
    out = DOCS / "Deep_Research_Streamlit_UI_Reference.docx"
    doc = Document()
    add_title_page(
        doc,
        "Deep Research — Streamlit Web UI Reference",
        f"Web application guide — {TODAY}",
    )

    add_heading(doc, "1. Launch", 1)
    add_para(doc, "From project root:", bold=True)
    add_bullets(
        doc,
        [
            "run_app.cmd — starts streamlit on app.py",
            "Default URL: http://localhost:8501",
            "Requires .env with CURSOR_API_KEY loaded at app startup",
        ],
    )

    add_heading(doc, "2. Layout", 1)
    add_bullets(
        doc,
        [
            "Left sidebar (30%): session list, search, status filter, pin/delete",
            "Main area (70%): New Research form or Session Detail view",
            "Wide layout with custom CSS (styles.py)",
        ],
    )

    add_heading(doc, "3. New Research Form", 1)
    add_bullets(
        doc,
        [
            "Topic (required), Field, Objectives (auto-derived if empty)",
            "Deliverables checkboxes — at least one required",
            "Link prior sessions (@context) — up to 3 completed sessions",
            "Depth, Audience",
            "Start Research — runs run_autonomous_mission() in st.status()",
        ],
    )

    add_heading(doc, "4. Session Detail View", 1)
    add_bullets(
        doc,
        [
            "Tabs: Answer, Papers, Report, Mind map, Slides, Activity",
            "Answer tab: turn timeline + follow-up form",
            "Papers tab: st.pdf() inline viewer + download",
            "Report tab: Markdown preview + DOCX download",
            "Mind map tab: PNG image viewer",
            "Slides tab: PPTX download",
            "Activity tab: persisted tool/status log",
        ],
    )

    add_heading(doc, "5. Session Actions", 1)
    add_table(
        doc,
        ["Action", "Behavior"],
        [
            ("Retry", "Re-run mission in same session folder"),
            ("Duplicate", "Prefill new form from session settings"),
            ("Branch", "New query with this session pre-linked as context"),
            ("Export ZIP", "Download all session files"),
            ("Delete", "Permanent removal of session folder"),
            ("Pin", "Float session to top of sidebar list"),
        ],
    )

    add_heading(doc, "6. Follow-Up Research", 1)
    add_para(
        doc,
        "On session detail Answer tab, enter a follow-up prompt with deliverable checkboxes. "
        "Calls run_follow_up_mission() which resumes cursor_agent_id if available, else "
        "reconstructs context from turns.json and linked sessions.",
    )

    add_heading(doc, "7. Sidebar Session Management", 1)
    add_bullets(
        doc,
        [
            "Search by topic, field, or tags",
            "Status filter: all, completed, error, running, pending, cancelled",
            "Clear failed — bulk delete all error sessions",
            "Bulk delete mode — multi-select then delete",
            "Delete confirmation dialog with checkbox",
        ],
    )

    doc.save(out)
    return out


def build_tools_reference() -> Path:
    out = DOCS / "Deep_Research_Custom_Tools_Reference.docx"
    doc = Document()
    add_title_page(
        doc,
        "Deep Research — Custom Tools Reference",
        f"Cursor SDK tool catalog — {TODAY}",
    )

    add_heading(doc, "1. Tool Registration", 1)
    add_para(
        doc,
        "Tools are registered in build_all_custom_tools() (deep_research/tools/__init__.py). "
        "Output tools come from cursor_agent_core.outputs. Domain tools stay in deep_research.",
    )

    add_heading(doc, "2. Always Available", 1)
    add_table(
        doc,
        ["Tool name", "Package", "Output"],
        [
            ("complete_mission", "deep_research", "completion.json manifest"),
            ("create_diagram", "cursor_agent_core", "session/Diagrams/ .mmd .png .svg"),
            ("create_slide_deck", "cursor_agent_core", "session/Slides/ .pptx"),
            ("create_research_report", "cursor_agent_core", "session/Reports/ .md .docx"),
        ],
    )

    add_heading(doc, "3. When TAVILY_API_KEY Set", 1)
    add_table(
        doc,
        ["Tool name", "Purpose"],
        [
            ("search_research_literature", "Tavily search with source_mode: academic|general|news"),
            ("download_research_pdfs", "Download PDFs from URLs to session/Papers/"),
        ],
    )

    add_heading(doc, "4. search_research_literature", 1)
    add_bullets(
        doc,
        [
            "source_mode=academic for peer-reviewed / institutional sources",
            "source_mode=general for cross-domain topics",
            "source_mode=news for policy, market, current events",
            "Returns URLs, snippets, and pdf_urls when available",
        ],
    )

    add_heading(doc, "5. create_diagram", 1)
    add_bullets(
        doc,
        [
            "Default diagram_type=mindmap when mindmap deliverable requested",
            "Accepts mermaid_source or branches JSON for mind maps",
            "Renders via mmdc (local node_modules or npx @mermaid-js/mermaid-cli)",
            "output_format: png, svg, or both",
        ],
    )

    add_heading(doc, "6. create_research_report", 1)
    add_bullets(
        doc,
        [
            "Structured sections: [{heading, body}, ...]",
            "Optional sources, key_takeaways, cited_papers paths",
            "Health disclaimer auto-added if not provided",
            "Produces paired .md and .docx files",
        ],
    )

    add_heading(doc, "7. complete_mission", 1)
    add_bullets(
        doc,
        [
            "Must be called last when mission finishes",
            "Requires summary; optional artifacts list, assumptions",
            "Writes completion.json in active session folder",
            "Agent instructed to stop after this tool succeeds",
        ],
    )

    doc.save(out)
    return out


def build_session_management_guide() -> Path:
    out = DOCS / "Deep_Research_Session_Management_Guide.docx"
    doc = Document()
    add_title_page(
        doc,
        "Deep Research — Session Management Guide",
        f"Per-query folders, linking, and follow-up — {TODAY}",
    )

    add_heading(doc, "1. Session ID Format", 1)
    add_para(doc, "{YYYYMMDD-HHMMSS}_{topic-slug}_{4-char-uuid}", bold=True)
    add_para(doc, "Example: 20260613-054121_crispr-gene-editing-in-agriculture_1f01")

    add_heading(doc, "2. meta.json Fields", 1)
    add_table(
        doc,
        ["Field", "Description"],
        [
            ("id, topic, field", "Identity and subject"),
            ("deliverables", "Requested output types"),
            ("status", "pending | running | completed | error | cancelled"),
            ("answer", "Latest agent text response"),
            ("agent_id", "Cursor SDK agent ID for resume"),
            ("linked_sessions", "Prior session IDs used as context"),
            ("parent_session_id", "Branch parent for session tree"),
            ("pinned, tags, notes", "User organization metadata"),
            ("num_turns, total_cost_usd", "Run statistics"),
        ],
    )

    add_heading(doc, "3. Linking Sessions (@context)", 1)
    add_para(
        doc,
        "When starting a new query, select up to 3 completed sessions to link. "
        "build_mission_prompt() injects prior answers, turn history, completion summaries, "
        "and artifact paths. Agent is instructed to build on prior findings without verbatim repeat.",
    )

    add_heading(doc, "4. Branching", 1)
    add_para(
        doc,
        "Branch from session detail sets prefill_linked to that session and opens New Research "
        "form with parent_session_id recorded in new session meta.json.",
    )

    add_heading(doc, "5. Multi-Turn Follow-Up", 1)
    add_numbered(
        doc,
        [
            "User submits follow-up prompt + deliverables on session detail page",
            "run_follow_up_mission() sets active session context",
            "Resumes agent via agent_id if stored in meta.json",
            "Appends turn to turns.json with prompt, answer, deliverables, timestamp",
            "New artifacts go into same session subfolders",
        ],
    )

    add_heading(doc, "6. Deleting Sessions", 1)
    add_bullets(
        doc,
        [
            "Single delete: sidebar Del button → confirmation dialog",
            "Bulk delete: toggle Bulk delete, select sessions, delete selected",
            "Clear failed: removes all sessions with status=error",
            "Delete is permanent (shutil.rmtree on session folder)",
        ],
    )

    add_heading(doc, "7. Export", 1)
    add_para(
        doc,
        "Export ZIP on session detail downloads meta.json, turns.json, activity.json, "
        "completion.json, and all artifacts in Papers/Reports/Diagrams/Slides/.",
    )

    doc.save(out)
    return out


def build_agents_master_update_note() -> Path:
    """Short addendum saved to parent docs referencing Deep Research."""
    out = AGENTS_DOCS / "Deep_Research_Agent_Workspace_Addendum.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    add_title_page(
        doc,
        "D:\\Agents — Deep Research Agent Addendum",
        f"Extends Agents_Folder_Master_Guide — {TODAY}",
    )
    add_heading(doc, "1. Third Agent Family Member", 1)
    add_para(
        doc,
        "Deep Research (D:\\Agents\\Deep research) is the first Cursor SDK agent. "
        "It uses cursor_agent_core, not agent_core, for runtime infrastructure.",
    )
    add_heading(doc, "2. Documentation Location", 1)
    add_bullets(
        doc,
        [
            "D:\\Agents\\Deep research\\docs\\ — Word reference documents",
            "D:\\Agents\\docs\\ARCHITECTURE.md — two-family model",
            "D:\\Agents\\docs\\CURSOR_SDK_AGENTS.md — scaffold new SDK agents",
        ],
    )
    add_heading(doc, "3. Quick Start", 1)
    add_para(doc, "cd D:\\Agents\\Deep research && run_app.cmd", bold=True)
    doc.save(out)
    return out


def main() -> int:
    builders = [
        build_master_guide,
        build_troubleshooting_guide,
        build_architecture_guide,
        build_cli_reference,
        build_streamlit_reference,
        build_tools_reference,
        build_session_management_guide,
        build_agents_master_update_note,
    ]
    for builder in builders:
        path = builder()
        print(f"Wrote {path}")
    return 0


if __name__ == "__main__":
    sys.path.insert(0, str(Path(__file__).resolve().parent))
    raise SystemExit(main())
